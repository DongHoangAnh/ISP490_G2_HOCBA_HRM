"""Dựng User manual (.docx) đúng khuôn bản Recruitment v1.0 của Việt.

Cách làm giống pipeline FS (`fs/gen_fs_docx.py`): mở bản Recruitment làm
"donor" để lấy nguyên styles.xml / numbering / theme / header-footer / sectPr
và các **bảng mẫu** (giữ đúng độ rộng cột), xoá sạch body rồi dựng lại theo
nội dung Python. Nhờ vậy hai user manual của nhóm đồng nhất tuyệt đối về
font, tiêu đề, kiểu bảng, canh lề và chú thích hình.

Dùng: from gen_um_docx import Doc  (xem build.py)
"""
import copy
import os

import docx
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Emu, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
DONOR = os.path.join(HERE, 'donor', 'UM-REC-v1.0.docx')

# Đo từ chính bản Recruitment: ảnh rộng 6.70", chú thích 9pt xám, in nghiêng.
IMG_WIDTH = Emu(6126480)
CAP_SIZE = Pt(9)
CAP_COLOR = RGBColor(0x80, 0x80, 0x80)
CELL_SZ = '21'  # half-points = 10.5pt, đúng cỡ chữ trong bảng của donor

# Bảng mẫu trong donor (chỉ số bảng → vai trò). Chọn theo VAI TRÒ chứ không
# theo số cột: hai bảng cùng 3 cột nhưng độ rộng cột khác nhau.
TPL_INDEX = {
    'history': 0,    # 6x6, hàng tiêu đề gộp ô "Change history"
    'sign': 1,       # 5x5, hàng tiêu đề gộp ô "FPT Signature"
    'toc': 2,        # 2x1, khung mục lục
    'process': 3,    # 3 cột — Process No | Definition | Note
    'stages': 4,     # 4 cột — # | Stage | Who | Where
    'roles': 5,      # 3 cột — Role | Can do | Cannot do
    'tabs': 6,       # 3 cột — Tab | Used for | Section
    'nav': 7,        # 3 cột — Screen | Path | Description
    'three': 13,     # 3 cột — bảng thường trong thân bài
    'two': 15,       # 2 cột — Tab | What it controls
    'faq': 16,       # 2 cột — Question | Answer
    'ref': 17,       # 2 cột — Document | Content
}

_TPL = None


def _templates():
    global _TPL
    if _TPL is None:
        d = docx.Document(DONOR)
        _TPL = {k: copy.deepcopy(d.tables[i]._tbl)
                for k, i in TPL_INDEX.items()}
    return _TPL


# ----------------------------------------------------------------------
# Thao tác XML mức thấp trên bảng mẫu
# ----------------------------------------------------------------------
def _rows(tbl):
    return tbl.findall(qn('w:tr'))


def _cells(tr):
    return tr.findall(qn('w:tc'))


def _set_cell(tc, text, bold):
    """Ghi text vào ô, giữ nguyên tcPr (độ rộng, nền, gộp ô)."""
    for p in tc.findall(qn('w:p')):
        tc.remove(p)
    for line in (text or '').split('\n') or ['']:
        p = tc.makeelement(qn('w:p'), {})
        tc.append(p)
        r = p.makeelement(qn('w:r'), {})
        rPr = r.makeelement(qn('w:rPr'), {})
        if bold:
            rPr.append(r.makeelement(qn('w:b'), {}))
        sz = r.makeelement(qn('w:sz'), {})
        sz.set(qn('w:val'), CELL_SZ)
        rPr.append(sz)
        r.append(rPr)
        t = r.makeelement(qn('w:t'), {})
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = line
        r.append(t)
        p.append(r)


def _make_table(kind, data, header=True):
    """Nhân bản hàng cuối của bảng mẫu cho đủ số hàng của `data`."""
    tbl = copy.deepcopy(_templates()[kind])
    trs = _rows(tbl)
    proto = copy.deepcopy(trs[-1])
    for tr in trs:
        tbl.remove(tr)
    for i, row in enumerate(data):
        tr = copy.deepcopy(proto)
        for j, tc in enumerate(_cells(tr)):
            _set_cell(tc, row[j] if j < len(row) else '', header and i == 0)
        tbl.append(tr)
    return tbl


def _fill_table(kind, cellmap):
    """Giữ NGUYÊN hình dạng bảng mẫu (kể cả hàng tiêu đề gộp ô), chỉ thay
    chữ ở các ô chỉ định: cellmap = {(row, col): (text, bold)}."""
    tbl = copy.deepcopy(_templates()[kind])
    trs = _rows(tbl)
    for (i, j), val in cellmap.items():
        text, bold = val if isinstance(val, tuple) else (val, False)
        tcs = _cells(trs[i])
        if j < len(tcs):
            _set_cell(tcs[j], text, bold)
    return tbl


def _toc_table():
    """Khung mục lục: giữ vỏ bảng của donor, thay ruột bằng field TOC MỚI
    (dirty=true → Word tự dựng lại khi mở / bấm Update Field). Không chép
    lại field cũ vì nó còn cache mục lục của module Tuyển dụng."""
    tbl = copy.deepcopy(_templates()['toc'])
    trs = _rows(tbl)
    _set_cell(_cells(trs[0])[0], 'Table of Contents', True)
    tc = _cells(trs[1])[0]
    for p in tc.findall(qn('w:p')):
        tc.remove(p)
    p = tc.makeelement(qn('w:p'), {})
    tc.append(p)
    fld = p.makeelement(qn('w:fldSimple'), {})
    fld.set(qn('w:instr'), r' TOC \o "1-3" \h \z \u ')
    fld.set(qn('w:dirty'), 'true')
    p.append(fld)
    inner = fld.makeelement(qn('w:p'), {})
    fld.append(inner)
    r = inner.makeelement(qn('w:r'), {})
    t = r.makeelement(qn('w:t'), {})
    t.text = ('Right-click here and choose "Update Field" to build '
              'the table of contents.')
    r.append(t)
    inner.append(r)
    return tbl


# ----------------------------------------------------------------------
# Tài liệu
# ----------------------------------------------------------------------
class Doc:
    def __init__(self):
        self.d = docx.Document(DONOR)
        body = self.d.element.body
        for ch in list(body.iterchildren()):
            if ch.tag != qn('w:sectPr'):
                body.remove(ch)
        self.sect = body.find(qn('w:sectPr'))
        self.fig = 0

    # --- đoạn văn ---
    def para(self, text='', bold=False, style=None, size=None, center=False,
             color=None, italic=False, page_break=False):
        p = self.d.add_paragraph(style=style)
        if center:
            p.alignment = 1
        if page_break:
            p.add_run().add_break(WD_BREAK.PAGE)
        if text:
            r = p.add_run(text)
            r.bold = bold or None
            r.italic = italic or None
            if size:
                r.font.size = size
            if color:
                r.font.color.rgb = color
        return p

    def h1(self, text):
        self.para(text, style='Heading 1')

    def h2(self, text):
        self.para(text, style='Heading 2')

    def h3(self, text):
        self.para(text, style='Heading 3')

    def bullets(self, items):
        for it in items:
            self.para(it, style='List Bullet')

    def note(self, text):
        """Dòng ghi chú nhỏ, xám — như dòng chú thích dưới mục lục của donor."""
        self.para(text, size=CAP_SIZE, color=CAP_COLOR)

    # --- bảng ---
    def table(self, kind, data, header=True):
        self.sect.addprevious(_make_table(kind, data, header))

    def raw_table(self, tbl):
        self.sect.addprevious(tbl)

    # --- hình + chú thích ---
    def figure(self, path, caption):
        """Ảnh căn giữa + dòng "Figure N – ..." 9pt nghiêng xám."""
        self.fig += 1
        p = self.d.add_paragraph()
        p.alignment = 1
        if os.path.exists(path):
            p.add_run().add_picture(path, width=IMG_WIDTH)
        else:
            raise FileNotFoundError(path)
        self.para(f'Figure {self.fig} – {caption}', center=True,
                  size=CAP_SIZE, color=CAP_COLOR, italic=True)

    def drop_unused_images(self):
        """Gỡ ảnh của donor còn sót trong package (~4.7MB ảnh chết)."""
        used = {r.get(qn('r:embed')) for r in
                self.d.element.body.findall('.//' + qn('a:blip'))}
        part = self.d.part
        for rid, rel in list(part.rels.items()):
            if rel.reltype.endswith('/image') and rid not in used:
                part.drop_rel(rid)

    def save(self, path):
        self.drop_unused_images()
        self.d.save(path)
        return path
