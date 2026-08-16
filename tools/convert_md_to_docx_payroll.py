import os
import re
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

if sys.platform == 'win32':
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except Exception:
        pass

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>\n'
        f'  <w:left w:val="none"/>\n'
        f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="A0A0A0"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>\n'
        f'  <w:insideV w:val="none"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def parse_markdown_table(lines):
    rows = []
    for line in lines:
        if re.match(r'^\s*\|?\s*:?-+:?\s*\|', line):
            continue  # Separator line
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    return rows

def md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = docx.Document()
    
    # Page Setup - Normal Margins (1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    lines = content.splitlines()
    i = 0
    in_code = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # Code block
        if line.strip().startswith('```'):
            if in_code:
                # Flush code block
                code_text = '\n'.join(code_lines)
                tbl = doc.add_table(rows=1, cols=1)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell = tbl.rows[0].cells[0]
                set_cell_background(cell, 'F4F5F7')
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x17, 0x2B, 0x4D)
                doc.add_paragraph()  # Spacing
                code_lines = []
                in_code = False
            else:
                in_code = True
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Table processing
        if line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            table_data = parse_markdown_table(table_lines)
            if table_data:
                cols_count = max(len(r) for r in table_data)
                tbl = doc.add_table(rows=len(table_data), cols=cols_count)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_borders(tbl)

                for r_idx, row in enumerate(table_data):
                    for c_idx, cell_text in enumerate(row):
                        if c_idx < cols_count:
                            cell = tbl.cell(r_idx, c_idx)
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_before = Pt(3)
                            p.paragraph_format.space_after = Pt(3)
                            if r_idx == 0:
                                set_cell_background(cell, '003366')  # Dark navy header
                                run = p.add_run(cell_text)
                                run.bold = True
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                run.font.size = Pt(10)
                            else:
                                if r_idx % 2 == 1:
                                    set_cell_background(cell, 'FFFFFF')
                                else:
                                    set_cell_background(cell, 'F8F9FA')  # Zebra striping
                                parse_formatted_text(p, cell_text)
                p_space = doc.add_paragraph()
                p_space.paragraph_format.space_before = Pt(0)
                p_space.paragraph_format.space_after = Pt(6)
            continue

        # Headings
        if line.startswith('# '):
            p = doc.add_heading(line[2:].strip(), level=1)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            for r in p.runs:
                r.font.name = 'Arial'
                r.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
                r.bold = True
        elif line.startswith('## '):
            p = doc.add_heading(line[3:].strip(), level=2)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            for r in p.runs:
                r.font.name = 'Arial'
                r.font.color.rgb = RGBColor(0x00, 0x40, 0x80)
                r.bold = True
        elif line.startswith('### '):
            p = doc.add_heading(line[4:].strip(), level=3)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            for r in p.runs:
                r.font.name = 'Calibri'
                r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
                r.bold = True
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:].strip(), level=4)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            for r in p.runs:
                r.font.name = 'Calibri'
                r.bold = True
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            parse_formatted_text(p, line[2:].strip())
        elif line.strip() == '---':
            # Divider line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/></w:pBdr>')
            p._p.get_or_add_pPr().append(pBdr)
        elif line.strip() == '':
            pass
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            parse_formatted_text(p, line.strip())

        i += 1

    doc.save(docx_path)
    print(f"Generated DOCX: {docx_path}")

def parse_formatted_text(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*|\`.*?\`|\*.*?\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x9C, 0x00, 0x06)
            rPr = run._r.get_or_add_rPr()
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
            rPr.append(shd)
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)

if __name__ == '__main__':
    md_dir = r'D:\FPT\DO_an\code\ISP490_G2_HOCBA_HRM\docs\specs\payroll'
    docx_dir = r'D:\FPT\DO_an\code\ISP490_G2_HOCBA_HRM\HRM_Học Bá\functional specs\funtional_specs_Payroll'

    mapping = [
        ('FS-PAY-001_Salary_Structure_Rule_Configuration_v1_0.md', 'FS-PAY-001_Salary_Structure_Rule_Configuration_v1.0.docx'),
        ('FS-PAY-002_Payslip_Computation_Engine_v1_0.md', 'FS-PAY-002_Payslip_Computation_Engine_v1.0.docx'),
        ('FS-PAY-003_Payslip_Lifecycle_Batch_Management_v1_0.md', 'FS-PAY-003_Payslip_Lifecycle_Batch_Management_v1.0.docx'),
        ('FS-PAY-004_Bank_Payment_File_Generation_v1_0.md', 'FS-PAY-004_Bank_Payment_File_Generation_v1.0.docx'),
        ('FS-PAY-005_Employee_Payslip_Confirmation_Email_v1_0.md', 'FS-PAY-005_Employee_Payslip_Confirmation_Email_v1.0.docx'),
    ]

    for md_file, docx_file in mapping:
        m_path = os.path.join(md_dir, md_file)
        d_path = os.path.join(docx_dir, docx_file)
        md_to_docx(m_path, d_path)
